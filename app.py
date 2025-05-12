from flask import Flask, request, jsonify, send_file, after_this_request
import os
import pandas as pd
from docx import Document
import uuid
import io
import zipfile
import requests
from pdfrw import PdfReader, PdfWriter, PageMerge
from reportlab.pdfgen import canvas
from reportlab.lib.pagesizes import letter
import subprocess
import time

app = Flask(__name__, static_folder='static', template_folder='static')
OUTPUT_FOLDER = 'output'
os.makedirs(OUTPUT_FOLDER, exist_ok=True)

current_session_id = None
current_file_prefix = None

COMPANY_TEMPLATES = {
    "ROYAL_SKY_INTERNATIONAL": 'templates/ROYAL',
    "NEW_VISION": 'templates/NEWVISION',
    "SNS_GLOBLE": 'templates/SNSGLOBLE'
}

SHEET_NAME = {
    "ROYAL_SKY_INTERNATIONAL": 'RS',
    "NEW_VISION": 'NV',
    "SNS_GLOBLE": 'SNS'
}

def replace_placeholders(doc, replacements):
    for paragraph in doc.paragraphs:
        for key, value in replacements.items():
            _replace_in_runs(paragraph.runs, f"{{{{{key}}}}}", str(value))
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    for key, value in replacements.items():
                        _replace_in_runs(paragraph.runs, f"{{{{{key}}}}}", str(value))

def _replace_in_runs(runs, placeholder, value):
    full_text = ''.join(run.text for run in runs)
    if placeholder not in full_text:
        return
    start = full_text.find(placeholder)
    while start != -1:
        end = start + len(placeholder)
        current = 0
        for run in runs:
            run_len = len(run.text)
            if current <= start < current + run_len:
                run_start = start - current
                run_end = min(run_len, end - current)
                before = run.text[:run_start]
                after = run.text[run_end:]
                run.text = before + value + after
                left = end - (current + run_len)
                if left > 0:
                    _remove_placeholder_from_next_runs(runs, runs.index(run)+1, left)
                break
            current += run_len
        full_text = ''.join(run.text for run in runs)
        start = full_text.find(placeholder)

def _remove_placeholder_from_next_runs(runs, start_idx, left):
    for i in range(start_idx, len(runs)):
        if left <= 0:
            break
        run = runs[i]
        if left >= len(run.text):
            left -= len(run.text)
            run.text = ''
        else:
            run.text = run.text[left:]
            left = 0

def fill_pdf_template(input_pdf_path, output_pdf_path, replacements):
    template_pdf = PdfReader(input_pdf_path)
    for page in template_pdf.pages:
        packet = io.BytesIO()
        can = canvas.Canvas(packet, pagesize=letter)
        y = 700
        for key, value in replacements.items():
            can.drawString(100, y, f"{key}: {value}")
            y -= 20
        can.save()
        packet.seek(0)
        overlay_pdf = PdfReader(packet)
        PageMerge(page).add(overlay_pdf.pages[0]).render()
    PdfWriter(output_pdf_path, trailer=template_pdf).write()

def convert_docx_to_pdf(docx_path, output_dir=None, timeout=30):
    """
    Convert a DOCX file to PDF using LibreOffice (soffice) on Linux.
    :param docx_path: Path to the input DOCX file.
    :param output_dir: Directory to save the output PDF. If None, uses the DOCX directory.
    :param timeout: Timeout for the conversion process in seconds.
    :return: Path to the generated PDF file.
    :raises: Exception if conversion fails.
    """
    if output_dir is None:
        output_dir = os.path.dirname(docx_path)
    os.makedirs(output_dir, exist_ok=True)
    cmd = [
        'soffice',  # or 'libreoffice' if that's the command on your system
        '--headless',
        '--convert-to', 'pdf',
        '--outdir', output_dir,
        docx_path
    ]
    try:
        result = subprocess.run(cmd, stdout=subprocess.PIPE, stderr=subprocess.PIPE, timeout=timeout)
        if result.returncode != 0:
            raise Exception(f"LibreOffice conversion failed: {result.stderr.decode()}")
        pdf_filename = os.path.splitext(os.path.basename(docx_path))[0] + '.pdf'
        pdf_path = os.path.join(output_dir, pdf_filename)
        if not os.path.exists(pdf_path):
            raise Exception("PDF file was not created.")
        return pdf_path
    except Exception as e:
        print(f"PDF conversion error: {e}")
        raise Exception(f"Error converting DOCX to PDF: {e}")

@app.route('/')
def index():
    return app.send_static_file('index.html')

@app.route('/set-template', methods=['POST'])
def set_template():
    data = request.get_json()
    dropdown_data = data.get("company_name_dropdown")
    TEMPLATE_FOLDER = COMPANY_TEMPLATES.get(dropdown_data)
    SHEET = SHEET_NAME.get(dropdown_data)
    app.config["TEMPLATE_FOLDER"] = TEMPLATE_FOLDER
    app.config["SHEET_NAME"] = SHEET
    google_sheet_url = "https://docs.google.com/spreadsheets/d/1vgXggucKcJ09xXJj-mjraFnk_PH3iCEKm1iv6Teq7UI/edit?gid=787616279#gid=787616279"
    app.config["GOOGLE_SHEET_URL"] = google_sheet_url
    print(f"Using folder: {TEMPLATE_FOLDER}")
    print(f"Using sheet: {SHEET}")
    print(f"Using sheet URL: {google_sheet_url}")
    return jsonify({"message": "Template and sheet set successfully"})

@app.route('/process', methods=['POST'])
def process():
    global current_session_id, current_file_prefix
    try:
        data = request.get_json()
        passport_number = data.get("passportNumber")
        output_format = data.get("outputFormat", "pdf")
        sheet_name = app.config.get("SHEET_NAME")
        google_sheet_url = app.config.get("GOOGLE_SHEET_URL")
        TEMPLATE_FOLDER = app.config.get("TEMPLATE_FOLDER")
        if not TEMPLATE_FOLDER or not os.path.exists(TEMPLATE_FOLDER):
            return jsonify({"success": False, "message": f"Template folder not found: {TEMPLATE_FOLDER}"})
        if not google_sheet_url:
            return jsonify({"success": False, "message": "Google Sheet URL not set. Please select a company first."})

        sheet_id = google_sheet_url.split("/d/")[1].split("/")[0]
        csv_url = f"https://docs.google.com/spreadsheets/d/{sheet_id}/gviz/tq?tqx=out:csv&sheet={sheet_name}"
        response = requests.get(csv_url)
        response.raise_for_status()
        csv_data = response.text
        lines = csv_data.splitlines()
        header_idx = None
        for i, line in enumerate(lines[:5]):
            if any('PASSPORTNO' in col.replace(' ', '').upper() for col in line.split(',')):
                header_idx = i
                break
        if header_idx is None:
            return jsonify({"success": False, "message": "PASSPORTNO column missing in sheet."})

        df = pd.read_csv(io.StringIO(csv_data), header=header_idx, dtype=str, skip_blank_lines=True)
        if 'PASSPORTNO' not in df.columns:
            return jsonify({"success": False, "message": "PASSPORTNO column missing in sheet."})

        df = df[df['PASSPORTNO'].notnull() & (df['PASSPORTNO'] != '')]
        passport_row = df[df['PASSPORTNO'].astype(str) == str(passport_number)]
        if passport_row.empty:
            return jsonify({"success": False, "message": "Passport number not found."})

        passport_data = passport_row.iloc[-1].copy()
        if 'VISAISSUEDATE' in passport_data and pd.notnull(passport_data['VISAISSUEDATE']):
            passport_data['VISAISSUEDATE'] = str(passport_data['VISAISSUEDATE'])

        for col in passport_data.index:
            val = passport_data[col]
            try:
                if isinstance(val, str) and val.isdigit():
                    passport_data[col] = int(val)
                elif isinstance(val, float) and val.is_integer():
                    passport_data[col] = int(val)
            except:
                pass

        country_name = passport_data['Country Name']
        sr_no = passport_data['srno']
        phoneno = passport_data['PHONENO']
        passport_data['PHONENO'] = phoneno

        templates_path = os.path.join(TEMPLATE_FOLDER, str(country_name))
        if not os.path.exists(templates_path):
            return jsonify({"success": False, "message": f"Templates not found for country: {country_name}"})

        session_id = str(uuid.uuid4())
        current_session_id = session_id
        current_file_prefix = f"{sr_no} {passport_number}"
        session_output = os.path.join(OUTPUT_FOLDER, session_id)
        os.makedirs(session_output, exist_ok=True)

        DOC_MAP = {
            'agreement': ('agreement.docx', 'Agreement'),
            'request_letter': ('request_letter.docx', 'Request Letter'),
            'afi_noc': ('afi_noc.docx', 'Affidavit')
        }
        selected_docs = data.get("selectedDocs", ['agreement', 'request_letter', 'afi_noc'])
        template_files = [DOC_MAP[key] for key in selected_docs if key in DOC_MAP]

        replacements = passport_data.to_dict()
        files = []

        for template_file, display_name in template_files:
            pdf_template_path = os.path.join(templates_path, template_file.replace('.docx', '.pdf'))
            output_name = f"{sr_no}-{display_name}"
            if os.path.exists(pdf_template_path):
                output_pdf = os.path.join(session_output, f"{output_name}.pdf")
                fill_pdf_template(pdf_template_path, output_pdf, replacements)
                files.append({
                    "name": f"{output_name}.pdf",
                    "url": f"/download/{session_id}/{output_name}.pdf"
                })
                continue

            template_path = os.path.join(templates_path, template_file)
            if not os.path.exists(template_path):
                continue
            doc = Document(template_path)
            replace_placeholders(doc, replacements)
            output_docx = os.path.join(session_output, f"{output_name}.docx")
            doc.save(output_docx)
            if output_format == "pdf":
                try:
                    output_pdf = convert_docx_to_pdf(output_docx, session_output)
                    files.append({
                        "name": f"{output_name}.pdf",
                        "url": f"/download/{session_id}/{output_name}.pdf"
                    })
                except Exception as e:
                    files.append({
                        "name": f"{output_name}.docx",
                        "url": f"/download/{session_id}/{output_name}.docx"
                    })
            else:
                files.append({
                    "name": f"{output_name}.docx",
                    "url": f"/download/{session_id}/{output_name}.docx"
                })

        return jsonify({"success": True, "files": files})

    except Exception as e:
        return jsonify({"success": False, "message": str(e)})

@app.route('/download/<session_id>/<filename>')
def download(session_id, filename):
    file_path = os.path.join(OUTPUT_FOLDER, session_id, filename)
    if not os.path.exists(file_path):
        return "File not found", 404

    @after_this_request
    def cleanup(response):
        try:
            # Wait a moment to ensure file is not locked (especially after conversion)
            time.sleep(0.3)
            os.remove(file_path)
            session_dir = os.path.join(OUTPUT_FOLDER, session_id)
            if os.path.exists(session_dir) and not os.listdir(session_dir):
                os.rmdir(session_dir)
        except Exception as e:
            print(f"Cleanup error: {e}")
        return response

    return send_file(file_path, as_attachment=True)

@app.route('/download-all', methods=['GET', 'POST'])
def download_all():
    global current_session_id, current_file_prefix
    if request.method == 'POST':
        data = request.get_json() or {}
        session_id = data.get('session_id', current_session_id)
        file_prefix = data.get('file_prefix', current_file_prefix)
    else:
        session_id = current_session_id
        file_prefix = current_file_prefix
    if not session_id:
        return "No files to download. Generate documents first.", 404
    session_dir = os.path.join(OUTPUT_FOLDER, session_id)
    if not os.path.exists(session_dir):
        return "Session files not found", 404
    memory_file = io.BytesIO()
    with zipfile.ZipFile(memory_file, 'w', zipfile.ZIP_DEFLATED) as zipf:
        for file in os.listdir(session_dir):
            file_path = os.path.join(session_dir, file)
            if os.path.isfile(file_path):
                zipf.write(file_path, arcname=file)
    memory_file.seek(0)
    zip_name = f"{file_prefix}.zip" if file_prefix else "all_documents.zip"

    @after_this_request
    def cleanup(response):
        import shutil
        try:
            shutil.rmtree(session_dir)
        except Exception as e:
            print(f"Error deleting session folder: {e}")
        return response

    return send_file(memory_file, mimetype='application/zip', as_attachment=True, download_name=zip_name)

if __name__ == '__main__':
    app.run(debug=False)
